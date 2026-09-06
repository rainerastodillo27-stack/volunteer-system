"""
Generate a professional Word Document (.docx) for White-Box Testing Documentation.
Matches academic thesis / capstone format (Table XX styling in Arial, clean headers, borders).
"""

import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_border(cell, **kwargs):
    """
    Set cell borders.
    kwargs: top, bottom, left, right
    values: dict(sz=12, val='single', color='000000')
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="{kwargs.get("top", {}).get("val", "none")}" w:sz="{kwargs.get("top", {}).get("sz", "0")}" w:space="0" w:color="{kwargs.get("top", {}).get("color", "auto")}"/>\n'
        f'  <w:left w:val="{kwargs.get("left", {}).get("val", "none")}" w:sz="{kwargs.get("left", {}).get("sz", "0")}" w:space="0" w:color="{kwargs.get("left", {}).get("color", "auto")}"/>\n'
        f'  <w:bottom w:val="{kwargs.get("bottom", {}).get("val", "none")}" w:sz="{kwargs.get("bottom", {}).get("sz", "0")}" w:space="0" w:color="{kwargs.get("bottom", {}).get("color", "auto")}"/>\n'
        f'  <w:right w:val="{kwargs.get("right", {}).get("val", "none")}" w:sz="{kwargs.get("right", {}).get("sz", "0")}" w:space="0" w:color="{kwargs.get("right", {}).get("color", "auto")}"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def create_whitebox_docx():
    doc = docx.Document()

    # Set page to Landscape for wide test tables
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.0)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    # Base styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9.5)
    font.color.rgb = RGBColor(30, 41, 59)

    # Document Header
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("WHITE-BOX TESTING DOCUMENTATION")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(21, 128, 61) # Green theme

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(12)
    run_sub = sub_p.add_run("System Implementation Verification, Path Testing & Execution Log Traceability\nNegrense Volunteers for Change (NVC Connect)")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = RGBColor(100, 116, 139)

    # Summary Stats Box
    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_table.autofit = False

    metrics = [
        ("TOTAL TEST CASES", "96"),
        ("PASSED", "96 (100%)"),
        ("FAILED", "0 (0%)"),
        ("COVERAGE RATE", "100%")
    ]
    col_widths = [Inches(2.4), Inches(2.4), Inches(2.4), Inches(2.4)]
    row = summary_table.rows[0]
    for idx, (label, val) in enumerate(metrics):
        cell = row.cells[idx]
        cell.width = col_widths[idx]
        set_cell_shading(cell, "F0FDF4")
        set_cell_border(cell, 
                        top=dict(val='single', sz='6', color='86EFAC'),
                        bottom=dict(val='single', sz='6', color='86EFAC'),
                        left=dict(val='single', sz='6', color='86EFAC'),
                        right=dict(val='single', sz='6', color='86EFAC'))
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        r_lbl = p.add_run(label + "\n")
        r_lbl.font.size = Pt(8)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = RGBColor(22, 101, 52)
        r_val = p.add_run(val)
        r_val.font.size = Pt(13)
        r_val.font.bold = True
        r_val.font.color.rgb = RGBColor(21, 128, 61)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Read markdown source
    md_path = os.path.join(os.getcwd(), "WHITE_BOX_TESTING.md")
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    # Split by Use Case sections
    use_case_pattern = re.compile(r'## (Use Case \d+: [^\n]+)\n\n(\|.+?)(?=\n---|\n# White-Box Testing Coverage|$)', re.DOTALL)
    matches = use_case_pattern.findall(content)

    table_number = 66 # Starting numbering like in the user's thesis image example

    for uc_title, table_md in matches:
        # Heading for Table
        t_header_p = doc.add_paragraph()
        t_header_p.paragraph_format.space_before = Pt(14)
        t_header_p.paragraph_format.space_after = Pt(2)
        t_header_p.paragraph_format.keep_with_next = True

        run_tbl_num = t_header_p.add_run(f"Table {table_number}\n")
        run_tbl_num.font.bold = True
        run_tbl_num.font.size = Pt(11)

        run_tbl_title = t_header_p.add_run(f"White-Box Testing – {uc_title.split(': ', 1)[1]}")
        run_tbl_title.font.size = Pt(11)
        run_tbl_title.font.color.rgb = RGBColor(15, 23, 42)

        # Parse table markdown
        lines = [line.strip() for line in table_md.strip().split('\n') if line.strip()]
        if len(lines) < 3:
            continue

        header_line = lines[0]
        headers = [c.strip() for c in header_line.split('|')[1:-1]]
        data_lines = lines[2:] # Skip separator line

        # Build table
        table = doc.add_table(rows=len(data_lines) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # Table Column Widths (Sum = ~9.8 inches)
        # ID (0.85"), Segment (1.15"), Desc (1.35"), Input (1.2"), Actual (2.4"), Expected (1.4"), Result (0.55"), Remarks (1.1")
        widths = [Inches(0.85), Inches(1.15), Inches(1.35), Inches(1.2), Inches(2.4), Inches(1.4), Inches(0.55), Inches(1.1)]

        # Header Row
        hdr_row = table.rows[0]
        hdr_row._tr.get_or_add_trPr().append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
        for idx, text in enumerate(headers):
            cell = hdr_row.cells[idx]
            cell.width = widths[idx]
            set_cell_shading(cell, "F8FAFC")
            set_cell_border(cell,
                            top=dict(val='single', sz='12', color='000000'), # 1.5pt solid top border
                            bottom=dict(val='single', sz='8', color='000000')) # 1.0pt bottom border
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 6) else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            r.font.bold = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(15, 23, 42)

        # Data Rows
        for r_idx, d_line in enumerate(data_lines):
            cols = [c.strip() for c in d_line.split('|')[1:-1]]
            row = table.rows[r_idx + 1]
            is_last_row = (r_idx == len(data_lines) - 1)

            for c_idx, val in enumerate(cols):
                cell = row.cells[c_idx]
                cell.width = widths[c_idx]
                
                # Bottom border on last row like formal tables
                if is_last_row:
                    set_cell_border(cell, bottom=dict(val='single', sz='12', color='000000'))

                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2.5)
                p.paragraph_format.space_after = Pt(2.5)
                p.paragraph_format.line_spacing = Pt(10.5)

                if c_idx == 0: # ID
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(val)
                    r.font.size = Pt(8)
                    r.font.bold = True
                elif c_idx == 6: # Result
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(val)
                    r.font.size = Pt(8.5)
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(22, 163, 74) # Green for Pass
                elif c_idx in (4, 5) and "[LOG:" in val: # Actual Behavior with log
                    parts = val.split("[LOG:", 1)
                    r1 = p.add_run(parts[0].strip() + "\n")
                    r1.font.size = Pt(8)
                    
                    log_text = "[LOG:" + parts[1].strip()
                    r2 = p.add_run(log_text)
                    r2.font.name = 'Consolas'
                    r2.font.size = Pt(7.5)
                    r2.font.color.rgb = RGBColor(3, 105, 161) # Cyan/Blue for log
                else:
                    r = p.add_run(val)
                    r.font.size = Pt(8)

        table_number += 1
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # White-Box Testing Coverage Section
    cov_p = doc.add_paragraph()
    cov_p.paragraph_format.space_before = Pt(20)
    cov_p.paragraph_format.space_after = Pt(4)
    run_cov_title = cov_p.add_run("White-Box Testing Coverage")
    run_cov_title.font.size = Pt(14)
    run_cov_title.font.bold = True
    run_cov_title.font.color.rgb = RGBColor(21, 128, 61)

    cov_intro = doc.add_paragraph(
        "The white-box testing focuses on the internal implementation of the system, including validation rules, "
        "conditional statements, authorization checks, database operations, status transitions, calculations, "
        "exception handling, and other execution paths."
    )
    cov_intro.paragraph_format.space_after = Pt(8)

    techniques = [
        ("Statement Coverage", "Verifies that the important executable statements within the tested modules are executed across normal and boundary conditions."),
        ("Branch Coverage", "Verifies both the true and false outcomes of decision statements such as validation checks, authorization checks, availability conditions, approval conditions, and status checks."),
        ("Condition Coverage", "Verifies the different conditions involved in system decisions, such as volunteer approval and availability, event capacity, proposal status, and project completion requirements."),
        ("Path Testing", "Verifies important successful and unsuccessful execution paths from the beginning of a function until its expected output."),
        ("Exception and Error Path Testing", "Verifies that database errors, invalid information, missing records, failed uploads, invalid status transitions, and similar exceptional conditions are safely handled.")
    ]

    for name, desc in techniques:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        r_name = p.add_run(f"• {name} – ")
        r_name.font.bold = True
        r_desc = p.add_run(desc)

    out_path = os.path.join(os.getcwd(), "White_Box_Testing_Documentation.docx")
    doc.save(out_path)
    print(f"[SUCCESS] Word document created at: {out_path}")

if __name__ == "__main__":
    create_whitebox_docx()
