"""
Remove redundant standalone 'Description: ...' paragraphs in all 4 Word test reports.
The description is already present inside each test case's table ('Description' column).
"""

import os
import docx

DOCS_DIR = os.path.abspath("docs/testings")

files_to_update = [
    ("Alpha_Blackbox_Test_Report.docx", 138),
    ("Beta_Blackbox_Test_Report.docx", 138),
    ("Alpha_Whitebox_Test_Report.docx", 100),
    ("Beta_Whitebox_Test_Report.docx", 100),
]

for fname, expected_removals in files_to_update:
    fpath = os.path.join(DOCS_DIR, fname)
    print(f"\nProcessing {fname}...")
    doc = docx.Document(fpath)

    initial_paragraphs = len(doc.paragraphs)
    initial_tables = len(doc.tables)
    initial_shapes = len(doc.inline_shapes)

    removed = 0
    for p in list(doc.paragraphs):
        if p.text.strip().startswith("Description:"):
            p._element.getparent().remove(p._element)
            removed += 1

    doc.save(fpath)
    print(f"  Removed {removed} 'Description:' paragraphs (expected: {expected_removals})")

    # Reload and verify
    doc2 = docx.Document(fpath)
    print(f"  Verified tables count: {len(doc2.tables)} (was {initial_tables})")
    print(f"  Verified images count: {len(doc2.inline_shapes)} (was {initial_shapes})")
    remaining = [p.text for p in doc2.paragraphs if p.text.strip().startswith("Description:")]
    print(f"  Remaining 'Description:' paragraphs: {len(remaining)}")

    if removed == expected_removals and len(remaining) == 0 and len(doc2.tables) == initial_tables and len(doc2.inline_shapes) == initial_shapes:
        print(f"  [SUCCESS] {fname} successfully cleaned and verified.")
    else:
        print(f"  [WARNING] Check {fname} for unexpected counts.")

print("\nAll 4 Word reports updated successfully.")
